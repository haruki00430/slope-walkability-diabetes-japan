"""
convert_to_jch_format.py
------------------------
Journal of Community Health (JCH) 投稿用 DOCX 変換スクリプト

【処理内容】
1. Vancouver 形式 "n)" の本文内引用 → JCH 形式 "[n]" に変換
2. 複合引用 "n), m)" → "[n, m]" に変換
3. 記号修正（Moran's / ΔAIC / α / β / ρ / R²）
4. 引用番号を初登場順に振り直し（旧35件→新34件、旧[22]削除）
5. References セクションを Declarations_JCH.md の新順序 APA-7 リストで自動置換
6. Declarations セクション（Funding / Ethics / Authors' Contributions など）を
   References の直前に正順で挿入
7. 出力: submission_package_JCH/Manuscript_JCH.docx

【使用方法】
    python convert_to_jch_format.py

【依存関係】
    pip install python-docx

【注意】
    - References の APA-7 変換元は Declarations_JCH.md（## References セクション）
    - 図表のキャプション内の引用も変換される
    - Zenodo DOI 確定後は Declarations_JCH.md の Data Availability を更新して再実行
"""

import re
import shutil
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("ERROR: python-docx が必要です。")
    print("       pip install python-docx  を実行してインストールしてください。")
    raise

# ============================================================
# パス設定
# ============================================================
SCRIPT_DIR = Path(__file__).parent
PROJECT_ROOT = SCRIPT_DIR.parent.parent  # NDB_XXX_slope_diabetes/

# 入力: バックアップ DOCX（元ファイルを保護）
INPUT_DOCX = SCRIPT_DIR / "Manuscript_JCH_backup_20260628.docx"

# 出力: JCH 形式 DOCX
OUTPUT_DOCX = SCRIPT_DIR / "Manuscript_JCH.docx"

# Declarations テキストファイル（Declarations_JCH.md の内容を参照）
DECLARATIONS_MD = SCRIPT_DIR / "Declarations_JCH.md"


# ============================================================
# 引用形式変換: "n)" → "[n]"
# ============================================================

# 本文内引用パターン（Vancouver 形式の数字引用）
#
# 【変換対象】
#   word4)         → word[4]       （単語直後）
#   methodology,4) → methodology,[4]（カンマ後）
#   I.11)          → I.[11]        （略語ピリオド後）
#   likelihood.9)  → likelihood.[9] （文末ピリオド後）
#   2020.10)       → 2020.[10]     （4桁年+ピリオド後）
#   8), 9)         → [8, 9]        （複合引用）
#
# 【変換対象外（偽陽性防止）】
#   0.05)  0.437)  0.001)  などの小数値・p値・係数
#   (N = 47)  (2020)  (2023) などの記述的括弧
#
# 識別ロジック:
#   単独引用は「直前が英字/句読点/閉じ括弧」の場合のみ変換。
#   小数値は前が数字 → 不一致でスキップ。
CITATION_PATTERNS = [
    # 4連続引用: "n), m), k), l)" → "[n, m, k, l]"
    (
        re.compile(r'(\d+)\),\s*(\d+)\),\s*(\d+)\),\s*(\d+)\)'),
        r'[\1, \2, \3, \4]'
    ),
    # 3連続引用: "n), m), k)" → "[n, m, k]"
    (
        re.compile(r'(\d+)\),\s*(\d+)\),\s*(\d+)\)'),
        r'[\1, \2, \3]'
    ),
    # 2連続引用: "n), m)" → "[n, m]"
    (
        re.compile(r'(\d+)\),\s*(\d+)\)'),
        r'[\1, \2]'
    ),
    # 範囲引用: "n)-m)" → "[n-m]"
    (
        re.compile(r'(\d+)\)-(\d+)\)'),
        r'[\1-\2]'
    ),
    # 単独引用①: 英字+ピリオド+引用番号 "I.11)" "likelihood.9)"
    #   lookbehind=英字、次に '.' を含む match → ".[\1]"
    (
        re.compile(r'(?<=[a-zA-Z])\.(\d{1,2})\)(?!\d)'),
        r'.[\1]'
    ),
    # 単独引用②: 4桁年+ピリオド+引用番号 "2020.10)"
    (
        re.compile(r'(?<=\d{4})\.(\d{1,2})\)(?!\d)'),
        r'.[\1]'
    ),
    # 単独引用③: 英字/カンマ/セミコロン/]/閉じ括弧の直後 "word4)" ",4)" ")2)"
    #   ※ 前が数字や小数点の場合は除外（p値・係数の保護）
    (
        re.compile(r'(?<=[a-zA-Z,;!?\]\)])(\d{1,2})\)(?!\d)'),
        r'[\1]'
    ),
]


# ============================================================
# 記号・表記の修正パターン（順序依存あり：長い文字列を先に）
# ============================================================

SYMBOL_REPLACEMENTS = [
    # Δ: 文字化け「?AIC」と表記不統一「Delta AIC」を統一
    ("Delta AIC",   "ΔAIC"),    # Delta AIC → ΔAIC（不統一解消）
    ("?AIC",        "ΔAIC"),    # ?AIC → ΔAIC（文字化け修正）
    # Moran’s / 所有格アポストロフィ文字化け
    ("Moran?s",     "Moran’s"),  # Moran?s → Moran’s
    ("outcome?s",   "outcome’s"),
    # ギリシャ文字（英語スペル → Unicode 記号）
    ("alpha = ",    "α = "),    # α =
    ("beta = ",     "β = "),    # β =
    ("rho = ",      "ρ = "),    # ρ =
    # R² （上付き 2 なし → Unicode 上付き文字）
    ("Pseudo R2 = ", "Pseudo R² = "),  # 長い方を先に
    ("R2 = ",        "R² = "),
    ("vs R2 =",      "vs R² ="),
    # Keywords: 旧8件 → TitlePage_JCH.md 指定の6件に修正（段落全体が1 run の場合に適用）
    ("Keywords: spatial epidemiology; spatial dependence; spatial structure; metabolic health; obesity; glycemic control; ecological study; NDB Open Data",
     "Keywords: built environment; walkability; diabetes; spatial autocorrelation; ecological study; Japan"),
]

# 正しい Keywords（TitlePage_JCH.md 準拠）
CORRECT_KEYWORDS = "Keywords: built environment; walkability; diabetes; spatial autocorrelation; ecological study; Japan"

# 著者情報（TitlePage_JCH.md 準拠）
AUTHOR_AFFILIATIONS = [
    "",
    "¹ Department of Epidemiology, Fukushima Medical University School of Medicine, Fukushima, Japan",
    "² Radiation Medical Science Center for the Fukushima Health Management Survey, Fukushima Medical University, Fukushima, Japan",
    "",
    "Corresponding Author: Haruki Saito",
    "Department of Epidemiology, Fukushima Medical University School of Medicine",
    "1 Hikarigaoka, Fukushima-shi, Fukushima 960-1295, Japan",
    "Email: m211039@fmu.ac.jp  |  ORCID: 0009-0009-7890-6068",
]


def apply_symbol_replacements_in_para(para) -> int:
    """記号置換を各 run に個別適用（書式を完全保持）。変換件数を返す。"""
    count = 0
    for run in para.runs:
        if not run.text:
            continue
        text = run.text
        for old, new in SYMBOL_REPLACEMENTS:
            if old in text:
                text = text.replace(old, new)
        if text != run.text:
            run.text = text
            count += 1
    return count


# ============================================================
# 引用番号の振り直し（Vancouver 初登場順に再番号付け）
# ============================================================
#
# 旧番号 → 新番号 のマッピング（Manuscript_JCH.docx 本文の初登場順）
# 旧 [22]（Nguyen et al.）は本文未引用のため削除。総数 34 件。
CITATION_RENUMBER_MAP = {
     1:  1,   2:  2,   3:  4,   4:  5,   5:  6,
     6:  7,   7:  8,   8:  9,   9: 10,  10: 11,
    11: 12,  12: 13,  13: 14,  14: 15,  15: 18,
    16: 22,  17: 23,  18: 31,  19: 24,  20: 34,
    21: 33,  23: 25,  24: 19,  25: 20,  26:  3,
    27: 16,  28: 21,  29: 29,  30: 30,  31: 17,
    32: 26,  33: 27,  34: 32,  35: 28,
}

_CITE_BLOCK_RE = re.compile(r'\[(\d[\d,\s\-]*)\]')


def _expand_citation_content(content: str) -> list:
    """'3, 4-7, 9' → [3, 4, 5, 6, 7, 9]"""
    nums = []
    for part in re.split(r',\s*', content.strip()):
        part = part.strip()
        m = re.match(r'^(\d+)-(\d+)$', part)
        if m:
            for n in range(int(m.group(1)), int(m.group(2)) + 1):
                nums.append(n)
        elif part.isdigit():
            nums.append(int(part))
    return nums


def _format_citation_list(nums: list) -> str:
    """[3, 4, 5, 7, 8] → '[3-5, 7, 8]'  (3以上連続 → 範囲、2連続 → 個別)"""
    if not nums:
        return ""
    sorted_nums = sorted(set(nums))
    groups = []
    i = 0
    while i < len(sorted_nums):
        start = sorted_nums[i]
        end = start
        while i + 1 < len(sorted_nums) and sorted_nums[i + 1] == end + 1:
            i += 1
            end = sorted_nums[i]
        if end - start >= 2:
            groups.append(f"{start}-{end}")
        elif end - start == 1:
            groups.append(str(start))
            groups.append(str(end))
        else:
            groups.append(str(start))
        i += 1
    return "[" + ", ".join(groups) + "]"


def renumber_citations_in_text(text: str) -> str:
    """本文中の [old_n] を [new_n] に振り直す（[n], [n-m], [n, m, k] 対応）。"""
    def _replace(match):
        old_nums = _expand_citation_content(match.group(1))
        new_nums = [CITATION_RENUMBER_MAP[n] for n in old_nums
                    if n in CITATION_RENUMBER_MAP]
        if not new_nums:
            return match.group(0)
        return _format_citation_list(new_nums)
    return _CITE_BLOCK_RE.sub(_replace, text)


def renumber_citations_in_para(para) -> int:
    """per-run で引用番号を振り直す（書式保持）。変換 run 数を返す。"""
    count = 0
    for run in para.runs:
        if not run.text:
            continue
        new_text = renumber_citations_in_text(run.text)
        if new_text != run.text:
            run.text = new_text
            count += 1
    return count


def convert_citations_in_text(text: str) -> str:
    """
    テキスト内のVancouver形式引用をJCH形式に変換する（最大3パス）。
    References セクション以降のテキストは変換しない（呼び出し元で制御）。
    5連続以上の引用は2パス目以降で残存を吸収する。
    """
    for _ in range(3):
        prev = text
        for pattern, replacement in CITATION_PATTERNS:
            text = pattern.sub(replacement, text)
        if text == prev:
            break  # 変更なければ早期終了
    return text


def is_references_section(para_text: str) -> bool:
    """References セクションの開始を検出する。"""
    stripped = para_text.strip()
    return stripped in ("References", "# References", "References\n")


# ============================================================
# Declarations テキストの生成
# ============================================================

DECLARATIONS_TEXT = """Statements and Declarations

Funding
No funding was received to assist with the preparation of this manuscript.

Competing Interests
The authors have no relevant financial or non-financial interests to disclose.

Ethics Approval
This study used publicly available, anonymized, aggregate-level data from the National Database of Health Insurance Claims and Specific Health Checkups (NDB) Open Data, released by the Ministry of Health, Labour and Welfare of Japan. No individual-level data were accessed. Ethics committee approval was not required under Japanese regulations for secondary analysis of de-identified aggregate data.

Consent to Participate
Not applicable (no individual participants were enrolled).

Consent for Publication
Not applicable.

Data Availability
The NDB Open Data (Ministry of Health, Labour and Welfare of Japan) used in this study are publicly available at https://www.mhlw.go.jp/stf/seisakunitsuite/bunya/0000177182.html. Digital elevation model data are available from the Geospatial Information Authority of Japan (https://www.gsi.go.jp/). Population Census data are available from e-Stat (https://www.e-stat.go.jp/). Analysis code is openly available at https://github.com/haruki00430/slope-walkability-diabetes-japan and archived on Zenodo at https://doi.org/10.5281/zenodo.20989085 (CC-BY 4.0).

Authors' Contributions
Haruki Saito conceptualized the study, performed data acquisition and statistical analysis, created all visualizations, and wrote the manuscript. Tetsuya Ohira reviewed the manuscript and provided critical revisions. All authors read and approved the final manuscript.
Conceptualization: Haruki Saito; Methodology: Haruki Saito; Formal analysis and investigation: Haruki Saito; Writing — original draft preparation: Haruki Saito; Writing — review and editing: Haruki Saito, Tetsuya Ohira; Resources: Haruki Saito.
"""


# ============================================================
# メイン変換処理
# ============================================================

def convert_docx(input_path: Path, output_path: Path) -> None:
    """DOCX を JCH 形式に変換して保存する。"""

    if not input_path.exists():
        raise FileNotFoundError(f"入力ファイルが見つかりません: {input_path}")

    print(f"入力: {input_path}")
    doc = Document(str(input_path))

    converted_count = 0
    symbol_count = 0
    renumber_count = 0
    in_references = False
    references_para_index = None
    references_para = None  # 段落オブジェクト（Pass 2/3 で再利用）

    # ---- Pass 1: 段落を走査し、引用変換・記号修正・番号振直 & References 位置を特定 ----
    for i, para in enumerate(doc.paragraphs):
        original_text = para.text

        # References セクション検出
        if is_references_section(original_text):
            in_references = True
            references_para_index = i
            references_para = para
            print(f"  [INFO] References セクション検出: 段落 {i}")
            continue

        # References 以降は変換しない
        if in_references:
            continue

        # 著者名段落: 上付き所属番号を付与
        if original_text.strip() == 'Haruki Saito':
            for run in para.runs:
                if 'Haruki Saito' in run.text:
                    run.text = run.text.replace('Haruki Saito', 'Haruki Saito¹')
            continue
        if original_text.strip() == 'Tetsuya Ohira':
            for run in para.runs:
                if 'Tetsuya Ohira' in run.text:
                    run.text = run.text.replace('Tetsuya Ohira', 'Tetsuya Ohira¹˒²')
            continue

        # Keywords 段落: 旧8件 → 正しい6件に一括置換
        if original_text.strip().startswith('Keywords:'):
            for run in para.runs:
                run.text = ''
            if para.runs:
                para.runs[0].text = CORRECT_KEYWORDS
            else:
                from docx.oxml import OxmlElement as _OE
                _r = _OE('w:r')
                _t = _OE('w:t')
                _t.text = CORRECT_KEYWORDS
                _r.append(_t)
                para._element.append(_r)
            continue

        # Step 1: 引用変換 n) → [n]（per-run、書式保持）
        new_text = convert_citations_in_text(original_text)
        if new_text != original_text:
            replace_text_in_para(para, original_text, new_text)
            converted_count += 1

        # Step 2: 記号・表記修正（per-run: Moran's / ΔAIC / α / β / ρ / R²）
        symbol_count += apply_symbol_replacements_in_para(para)

        # Step 3: 引用番号振り直し [old] → [new]（per-run、書式保持）
        renumber_count += renumber_citations_in_para(para)

    print(f"  [INFO] 引用変換 n)→[n]: {converted_count} 箇所")
    print(f"  [INFO] 記号・表記修正: {symbol_count} run")
    print(f"  [INFO] 引用番号振り直し: {renumber_count} run")

    # ---- Pass 2: References セクションを Declarations_JCH.md の APA-7 リストで置き換え ----
    if references_para is not None:
        ref_count = replace_references_section(doc, references_para, DECLARATIONS_MD)
        print(f"  [INFO] References 置き換え完了（{ref_count} 件）")
    else:
        print(f"  [WARNING] References セクションが見つかりませんでした")

    # ---- Pass 3: 著者所属・対応著者情報を著者名の後に挿入 ----
    insert_author_affiliations(doc)

    # ---- Pass 4: Declarations セクションを References の前に挿入 ----
    if references_para is not None:
        insert_declarations_before_references(doc, references_para)
        print(f"  [INFO] Declarations セクションを挿入しました")
    else:
        print(f"  [WARNING] Declarations を末尾に追加します（フォールバック）")
        doc.add_paragraph()
        doc.add_heading("Statements and Declarations", level=1)
        for line in DECLARATIONS_TEXT.strip().split("\n")[1:]:
            doc.add_paragraph(line)

    # ---- 保存 ----
    doc.save(str(output_path))
    print(f"出力: {output_path}")
    print("変換完了。")


def replace_text_in_para(para, original_text: str, new_text: str) -> None:
    """
    段落内の各 run を個別に引用変換する（書式を完全保持）。

    各 run のテキストに convert_citations_in_text を直接適用する。
    run を結合しないため bold / italic / font size / color などが維持される。
    引用番号が複数 run にまたがる場合は変換不可だが、
    Vancouver 形式 n) は通常 1 run 内に完結するため実用上問題なし。
    """
    if not para.runs:
        return

    for run in para.runs:
        if not run.text:
            continue
        converted = convert_citations_in_text(run.text)
        if converted != run.text:
            run.text = converted


def insert_declarations_before_references(doc: Document, ref_para) -> None:
    """
    References 段落の直前に Declarations セクションを正順で挿入する。

    【挿入ロジック】
    parent.insert(idx, elem) は idx 位置に elem を挿入し、
    それ以降の要素を右にシフトする。
    「ref_elem の直前に insert」を前から順に繰り返すと、
    最初に挿入した要素が ref から最も遠い位置（= セクション先頭）に確定し、
    後続の要素がそれに続いて正順に並ぶ。

    旧実装の reversed() バグ（逆順挿入）と見出し二重化を修正済み。
    """
    from docx.oxml import OxmlElement

    ref_elem = ref_para._element
    parent = ref_elem.getparent()

    # DECLARATIONS_TEXT の1行目 "Statements and Declarations" は
    # Heading1 スタイルで別途処理するため content_lines は2行目以降
    all_lines = DECLARATIONS_TEXT.strip().split("\n")
    assert all_lines[0] == "Statements and Declarations", (
        f"DECLARATIONS_TEXT の先頭行が不正: {all_lines[0]!r}"
    )
    content_lines = all_lines[1:]

    # Step 1: Heading1 段落を ref の直前に挿入（= セクション先頭になる）
    heading_para = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), 'Heading1')
    pPr.append(pStyle)
    heading_para.append(pPr)
    h_run = OxmlElement('w:r')
    h_text = OxmlElement('w:t')
    h_text.text = "Statements and Declarations"
    h_run.append(h_text)
    heading_para.append(h_run)
    parent.insert(list(parent).index(ref_elem), heading_para)

    # Step 2: コンテンツ行を前から順に ref の直前へ挿入
    # 毎回 ref の直前（= 常に末尾直前）に追加されるため正順で並ぶ
    for line in content_lines:
        new_para = OxmlElement('w:p')
        if line:
            c_run = OxmlElement('w:r')
            c_text = OxmlElement('w:t')
            c_text.text = line
            c_text.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            c_run.append(c_text)
            new_para.append(c_run)
        parent.insert(list(parent).index(ref_elem), new_para)

    # Step 3: 見出しの直前に空白段落を挿入（直前コンテンツとの余白）
    blank = OxmlElement('w:p')
    parent.insert(list(parent).index(heading_para), blank)


def parse_references_from_md(md_path: Path) -> list:
    """Declarations_JCH.md の ## References セクションから APA-7 エントリを解析する。

    Returns:
        list of str: "[1] Author ... *Journal*, *vol* ..." 形式（*...*の italic マーク保持）
    """
    text = md_path.read_text(encoding='utf-8')

    # ## References セクションを抽出（次の ## またはファイル末尾まで）
    m = re.search(r'^## References\s*\n(.*?)(?:^## |\Z)', text, re.DOTALL | re.MULTILINE)
    if not m:
        return []

    ref_section = m.group(1)
    entries = []
    for line in ref_section.split('\n'):
        line = line.strip()
        if re.match(r'^\[\d+\] ', line):
            # *...* マークはそのまま保持（replace_references_section で italic run に変換）
            entries.append(line)

    return entries


def _build_ref_para(entry_with_marks: str):
    """APA-7 参考文献テキスト（*...* マーク付き）を italic run 構造で構築した段落要素を返す。

    *...*  で囲まれた部分（誌名・巻号）を w:i（italic）run として出力する。
    """
    from docx.oxml import OxmlElement

    new_para = OxmlElement('w:p')
    # *...* でテキストを分割: 偶数インデックス=plain、奇数インデックス=italic
    parts = re.split(r'\*([^*]+)\*', entry_with_marks)
    for idx, part in enumerate(parts):
        if not part:
            continue
        run = OxmlElement('w:r')
        if idx % 2 == 1:  # italic segment
            rPr = OxmlElement('w:rPr')
            italic = OxmlElement('w:i')
            italic_cs = OxmlElement('w:iCs')
            rPr.append(italic)
            rPr.append(italic_cs)
            run.append(rPr)
        text_elem = OxmlElement('w:t')
        text_elem.text = part
        text_elem.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        run.append(text_elem)
        new_para.append(run)
    return new_para


def replace_references_section(doc: Document, ref_para, md_path: Path) -> int:
    """DOCX の References 見出し以降の旧参考文献を削除し、
    Declarations_JCH.md の新順序 APA-7 リスト（italic 付き）で置き換える。

    Returns:
        int: 挿入した参考文献件数
    """
    entries = parse_references_from_md(md_path)
    if not entries:
        print("  [WARNING] 参考文献が解析できませんでした（Declarations_JCH.md を確認してください）")
        return 0

    ref_elem = ref_para._element
    parent = ref_elem.getparent()

    # ref_elem 以降の兄弟要素を走査し、旧参考文献段落を収集・削除
    # "Tables" / "Figures" / "Figure" で始まる段落に達したら停止
    to_remove = []
    found = False
    for child in list(parent):
        if found:
            child_text = ''.join(
                t.text or ''
                for t in child.iter(
                    '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'
                )
            ).strip()
            if child_text.startswith(('Tables', 'Figures', 'Figure', 'Supplementary')):
                break
            to_remove.append(child)
        elif child is ref_elem:
            found = True

    for elem in to_remove:
        parent.remove(elem)

    print(f"  [INFO] 旧 References 段落 {len(to_remove)} 件を削除")

    # 新参考文献を ref_elem の直後に正順で挿入（reversed で addnext → 正順）
    for entry in reversed(entries):
        new_para = _build_ref_para(entry)
        ref_elem.addnext(new_para)

    return len(entries)


def insert_author_affiliations(doc: Document) -> None:
    """Tetsuya Ohira 段落の後に所属・対応著者情報を挿入する。"""
    from docx.oxml import OxmlElement

    ohira_para = None
    for para in doc.paragraphs:
        if para.text.strip() == 'Tetsuya Ohira¹˒²':
            ohira_para = para
            break
    # 上付き未付与の場合（バックアップ DOCX からそのまま Pass 1 前に呼ばれた等）もフォールバック
    if ohira_para is None:
        for para in doc.paragraphs:
            if 'Tetsuya Ohira' in para.text:
                ohira_para = para
                break

    if ohira_para is None:
        print("  [WARNING] Tetsuya Ohira の段落が見つかりません（著者所属をスキップ）")
        return

    ohira_elem = ohira_para._element

    # addnext に reversed で挿入 → 正順で並ぶ
    for line in reversed(AUTHOR_AFFILIATIONS):
        new_para = OxmlElement('w:p')
        if line:
            run = OxmlElement('w:r')
            text = OxmlElement('w:t')
            text.text = line
            text.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            run.append(text)
            new_para.append(run)
        ohira_elem.addnext(new_para)

    print(f"  [INFO] 著者所属・対応著者情報を挿入しました（{len(AUTHOR_AFFILIATIONS)} 行）")


# ============================================================
# エントリポイント
# ============================================================

if __name__ == "__main__":
    print("=" * 60)
    print("JCH Format Converter")
    print("Vancouver 1) → JCH [1] 変換")
    print("=" * 60)
    print()

    # 出力ファイルが既に存在する場合の確認
    if OUTPUT_DOCX.exists():
        print(f"注意: {OUTPUT_DOCX.name} が既に存在します。上書きします。")

    convert_docx(INPUT_DOCX, OUTPUT_DOCX)

    print()
    print("=" * 60)
    print("次のステップ:")
    print("1. Manuscript_JCH.docx を Word で開いて内容を確認")
    print("   - Declarations が References の前に正順で挿入されているか")
    print("   - References が [1]〜[34] の APA-7 形式（34件）になっているか")
    print("   - 引用番号が本文内で [1]〜[34] になっているか")
    print("2. Keywords を 6 件（現在）以内に確認（JCH 上限 6 件）")
    print("3. 数式変数（para 58, 59 付近）を Word で手動補完")
    print("4. Zenodo DOI 取得後、Declarations_JCH.md の Data Availability を更新して再実行")
    print("5. 著者名・所属のプレースホルダーを実名に置き換え")
    print("6. Editorial Manager にアップロード")
    print("=" * 60)
