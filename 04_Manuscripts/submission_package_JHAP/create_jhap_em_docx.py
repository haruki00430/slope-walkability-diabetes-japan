# -*- coding: utf-8 -*-
"""
Health & Place (JHAP) 投稿用 DOCX を準備する。

方針: ORCP パッケージの正本をコピーし、投稿規定・転投稿で
必須の差分だけを適用する。それ以外の文言は変更しない。
"""

from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph

PACKAGE_DIR = Path(__file__).resolve().parent
ORCP_DIR = PACKAGE_DIR.parent / "submission_package_ORCP"

# ORCP → JHAP: ファイル名のみ変更してそのままコピー
COPY_AS_IS = [
    ("TitlePage_ORCP.docx", "TitlePage_JHAP.docx"),
    ("Highlights_ORCP.docx", "Highlights_JHAP.docx"),
    ("Declaration_of_Interest_ORCP.docx", "Declaration_of_Interest_JHAP.docx"),
    ("Ethical_Statement_ORCP.docx", "Ethical_Statement_JHAP.docx"),
    ("STROBE_Checklist_ORCP.docx", "STROBE_Checklist_JHAP.docx"),
]

# カバーレター: 必須差分のみ（誌名・適合理由・投稿履歴・日付）
COVER_REPLACEMENTS = [
    (
        "for consideration for publication in Obesity Research & "
        "Clinical Practice as an Original Research Article.",
        "for consideration for publication in Health & Place as an "
        "original research article.",
    ),
    (
        "Obesity Research & Clinical Practice is ideally suited for this "
        "work. The journal's focus on the epidemiology and complication of "
        "obesity, with particular welcome for Asia Oceania studies, "
        "directly matches our prefecture-level ecological analysis of "
        "obesity and glycemic control across Japan.",
        "Health & Place is ideally suited for this work. The journal's "
        "interdisciplinary focus on how place shapes health and "
        "health-related experiences directly matches our prefecture-level "
        "spatial epidemiology study across Japan.",
    ),
    (
        "It was previously submitted to the Journal of Community Health "
        "(Manuscript ID: JOHE-D-26-01791) and was declined after "
        "editorial review on 30 June 2026. It is not currently under "
        "consideration at any other journal.",
        "It was previously submitted to the Journal of Community Health "
        "(Manuscript ID: JOHE-D-26-01791; declined after editorial "
        "review on 30 June 2026) and to Obesity Research & Clinical "
        "Practice (Manuscript ID: ORCP-D-26-00394; declined after "
        "editorial review on 7 July 2026). It is not currently under "
        "consideration at any other journal.",
    ),
]


def set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    """段落テキストを置換する。"""
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def copy_unchanged_files() -> None:
    """ORCP 正本をファイル名変更のみでコピーする。"""
    for src_name, dst_name in COPY_AS_IS:
        src = ORCP_DIR / src_name
        dst = PACKAGE_DIR / dst_name
        if not src.exists():
            raise FileNotFoundError(src)
        shutil.copy2(src, dst)
        print(f"Copied (unchanged): {dst_name}")


def patch_submission_dates() -> None:
    """宣言類 DOCX の日付行のみ 7 July 2026 に更新する。"""
    targets = [
        "Declaration_of_Interest_JHAP.docx",
        "Ethical_Statement_JHAP.docx",
        "STROBE_Checklist_JHAP.docx",
    ]
    for name in targets:
        path = PACKAGE_DIR / name
        doc = Document(str(path))
        for paragraph in doc.paragraphs:
            if paragraph.text.strip() == "Date: 30 June 2026":
                set_paragraph_text(paragraph, "Date: 7 July 2026")
        doc.save(str(path))
        print(f"Patched dates: {name}")


def create_cover_letter() -> None:
    """ORCP カバーレターをコピーし、必須差分のみ適用する。"""
    src = ORCP_DIR / "CoverLetter_ORCP.docx"
    dst = PACKAGE_DIR / "CoverLetter_JHAP.docx"
    if not src.exists():
        raise FileNotFoundError(src)
    shutil.copy2(src, dst)

    doc = Document(str(dst))
    for paragraph in doc.paragraphs:
        text = paragraph.text
        updated = text
        for old, new in COVER_REPLACEMENTS:
            updated = updated.replace(old, new)
        if updated != text:
            set_paragraph_text(paragraph, updated)

    # 日付行のみ更新（JCH 謝絶日は変更しない）
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == "30 June 2026":
            set_paragraph_text(paragraph, "7 July 2026")
            break

    doc.save(str(dst))
    print("Patched (required diffs only): CoverLetter_JHAP.docx")


def main() -> None:
    copy_unchanged_files()
    patch_submission_dates()
    create_cover_letter()
    print("JHAP EM DOCX ready (minimal-change policy)")


if __name__ == "__main__":
    main()
