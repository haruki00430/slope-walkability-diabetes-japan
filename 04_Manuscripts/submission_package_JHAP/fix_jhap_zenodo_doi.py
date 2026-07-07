# -*- coding: utf-8 -*-
"""Data availability / Title Page の Zenodo DOI を v1.0.1 に更新する。"""

from __future__ import annotations

from pathlib import Path

from docx import Document

PACKAGE_DIR = Path(__file__).resolve().parent

OLD_DOI = "10.5281/zenodo.20989085"
NEW_DOI = "10.5281/zenodo.21237968"

TARGETS = [
    PACKAGE_DIR / "Manuscript_JHAP_anonymized.docx",
    PACKAGE_DIR / "TitlePage_JHAP.docx",
]


def replace_in_paragraph(paragraph, old: str, new: str) -> bool:
    """段落内の DOI を置換する。"""
    if old not in paragraph.text:
        return False
    if paragraph.runs:
        paragraph.runs[0].text = paragraph.text.replace(old, new)
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(paragraph.text.replace(old, new))
    return True


def replace_in_doc(path: Path) -> int:
    """DOCX 内の Zenodo DOI を置換し、置換回数を返す。"""
    doc = Document(path)
    count = 0
    for paragraph in doc.paragraphs:
        if replace_in_paragraph(paragraph, OLD_DOI, NEW_DOI):
            count += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if replace_in_paragraph(paragraph, OLD_DOI, NEW_DOI):
                        count += 1
    if count:
        doc.save(path)
    return count


def main() -> None:
    for path in TARGETS:
        if not path.exists():
            raise FileNotFoundError(path)
        n = replace_in_doc(path)
        print(f"{path.name}: {n} replacement(s)")


if __name__ == "__main__":
    main()
