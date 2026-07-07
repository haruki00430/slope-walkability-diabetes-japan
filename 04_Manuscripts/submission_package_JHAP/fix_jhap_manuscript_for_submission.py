# -*- coding: utf-8 -*-
"""
Health & Place 投稿規定で必須の修正のみを匿名原稿に適用する。

方針: ORCP 正本（Manuscript_ORCP_anonymized.docx）をベースに、
規定上追加が必要な箇所だけを変更する。それ以外は一切変更しない。
"""

from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph

PACKAGE_DIR = Path(__file__).resolve().parent
ORCP_DIR = PACKAGE_DIR.parent / "submission_package_ORCP"
SOURCE = ORCP_DIR / "Manuscript_ORCP_anonymized.docx"
MANUSCRIPT = PACKAGE_DIR / "Manuscript_JHAP_anonymized.docx"

# Title Page の Data Availability と同一文言（新規作文を避ける）
DATA_AVAILABILITY_HEADING = "Data availability"
DATA_AVAILABILITY_TEXT = (
    "The NDB Open Data (Ministry of Health, Labour and Welfare of Japan) used in this "
    "study are publicly available at "
    "https://www.mhlw.go.jp/stf/seisakunitsuite/bunya/0000177182.html. Digital elevation "
    "model data are available from the Geospatial Information Authority of Japan "
    "(https://www.gsi.go.jp/). Population Census data are available from e-Stat "
    "(https://www.e-stat.go.jp/). Analysis code is openly available at "
    "https://github.com/haruki00430/slope-walkability-diabetes-japan and archived on "
    "Zenodo (https://doi.org/10.5281/zenodo.21237968)."
)

MAP_NOTE = (
    "Map lines delineate study areas and do not necessarily depict accepted national "
    "boundaries."
)


def set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    """段落テキストをスタイル維持のまま置換する。"""
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def restore_from_orcp() -> None:
    """ORCP 匿名原稿正本からコピーして上書きする。"""
    if not SOURCE.exists():
        raise FileNotFoundError(f"ORCP source missing: {SOURCE}")
    shutil.copy2(SOURCE, MANUSCRIPT)
    print(f"Restored from {SOURCE.name}")


def add_data_availability_section(doc: Document) -> None:
    """H&P 必須: AI 開示直前に Data availability を追加（既存ならスキップ）。"""
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(DATA_AVAILABILITY_HEADING):
            print("Data availability: already present, skip")
            return

    ai_para = next(
        p
        for p in doc.paragraphs
        if p.text.strip().startswith("Declaration of generative AI")
    )
    heading = ai_para.insert_paragraph_before(DATA_AVAILABILITY_HEADING)
    heading.style = doc.styles["Heading 1"]
    body = ai_para.insert_paragraph_before(DATA_AVAILABILITY_TEXT)
    body.style = doc.styles["Normal"]
    print("Data availability: added")


def append_map_note_to_lisa_caption(doc: Document) -> None:
    """H&P 必須: LISA 地図キャプションに境界注記を1箇所だけ追加。"""
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if "LISA cluster map" in text and MAP_NOTE not in text:
            set_paragraph_text(paragraph, f"{text} {MAP_NOTE}")
            print("Map note: added to LISA caption (body)")
            return

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = paragraph.text.strip()
                    if "LISA cluster map" in text and MAP_NOTE not in text:
                        set_paragraph_text(paragraph, f"{text} {MAP_NOTE}")
                        print("Map note: added to LISA caption (table)")
                        return

    print("Map note: LISA caption not found (no change)")


def verify_anonymization(doc: Document) -> None:
    """著者識別情報が含まれていないか簡易チェックする。"""
    banned = [
        "Haruki Saito",
        "Tetsuya Ohira",
        "Fukushima Medical University",
        "m211039@fmu.ac.jp",
    ]
    full_text = "\n".join(p.text for p in doc.paragraphs)
    for token in banned:
        if token in full_text:
            raise ValueError(f"Anonymization breach: found '{token}'")


def main() -> None:
    restore_from_orcp()
    doc = Document(MANUSCRIPT)
    add_data_availability_section(doc)
    append_map_note_to_lisa_caption(doc)
    verify_anonymization(doc)
    doc.save(MANUSCRIPT)
    print(f"Saved {MANUSCRIPT.name} (required H&P edits only)")


if __name__ == "__main__":
    main()
