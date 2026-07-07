# -*- coding: utf-8 -*-
"""Fig. 3 埋め込みキャプションに地図境界注記を追加する（本文は変更しない）。"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.text.paragraph import Paragraph

PACKAGE_DIR = Path(__file__).resolve().parent
MANUSCRIPT = PACKAGE_DIR / "Manuscript_JHAP_anonymized.docx"

MAP_NOTE = (
    "Map lines delineate study areas and do not necessarily depict "
    "accepted national boundaries."
)


def set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    """段落テキストをスタイル維持のまま置換する。"""
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def append_map_note_if_missing(text: str) -> str | None:
    """注記が未追加なら付与した文字列を返す。"""
    if MAP_NOTE in text:
        return None
    return f"{text.rstrip()} {MAP_NOTE}"


def fix_embedded_fig3_caption(doc: DocxDocument) -> bool:
    """図テーブル内の Fig. 3 キャプションに地図注記を追加する。"""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = paragraph.text.strip()
                    if not text.startswith("Fig. 3:"):
                        continue
                    updated = append_map_note_if_missing(text)
                    if updated:
                        set_paragraph_text(paragraph, updated)
                        print("Map note: added to embedded Fig. 3 caption")
                        return True
    return False


def main() -> None:
    doc = Document(str(MANUSCRIPT))
    if fix_embedded_fig3_caption(doc):
        doc.save(str(MANUSCRIPT))
        print(f"Saved {MANUSCRIPT.name}")
    else:
        print("Embedded Fig. 3 caption already has map note or not found")


if __name__ == "__main__":
    main()
